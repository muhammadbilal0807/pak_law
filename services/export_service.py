# services/export_service.py (UPDATED)
import io
from docx import Document
from fpdf import FPDF

def export_txt(content: str) -> bytes:
    return content.encode('utf-8')

def export_docx(content: str) -> bytes:
    doc = Document()
    doc.add_heading('Pak Law AI - Export', 0)
    doc.add_paragraph(content)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def export_pdf(content: str) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    # Basic cleanup for PDF
    safe_content = content.encode('latin-1', 'replace').decode('latin-1')
    pdf.multi_cell(0, 10, txt=safe_content)
    # FIXED: Use proper output method
    return pdf.output(dest='S').encode('latin-1')